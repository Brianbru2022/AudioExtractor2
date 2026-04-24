from __future__ import annotations

import csv
import json
import os
import re
from copy import copy
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from app.core.config import config
from app.repositories.extraction import (
    ExtractionEntityRepository,
    ExtractionEvidenceRepository,
    ExtractionRunRepository,
    ExtractionSummaryRepository,
)
from app.repositories.exports import ExportRunRepository
from app.repositories.meetings import ArtifactRepository, MeetingRepository, RunRepository, SourceFileRepository
from app.repositories.transcription import TranscriptSegmentRepository, TranscriptionRunRepository
from app.services.exports.models import ExportDescriptor, ExportFormat, ExportOptions, ExportProfile
from app.utils.files import ensure_directory, utc_now_iso


class ExportService:
    def __init__(
        self,
        *,
        meeting_repository: MeetingRepository,
        source_file_repository: SourceFileRepository,
        preprocessing_run_repository: RunRepository,
        transcription_run_repository: TranscriptionRunRepository,
        transcript_segment_repository: TranscriptSegmentRepository,
        extraction_run_repository: ExtractionRunRepository,
        action_repository: ExtractionEntityRepository,
        decision_repository: ExtractionEntityRepository,
        risk_repository: ExtractionEntityRepository,
        question_repository: ExtractionEntityRepository,
        topic_repository: ExtractionEntityRepository,
        evidence_repository: ExtractionEvidenceRepository,
        summary_repository: ExtractionSummaryRepository,
        artifact_repository: ArtifactRepository,
        export_run_repository: ExportRunRepository,
    ) -> None:
        self.meeting_repository = meeting_repository
        self.source_file_repository = source_file_repository
        self.preprocessing_run_repository = preprocessing_run_repository
        self.transcription_run_repository = transcription_run_repository
        self.transcript_segment_repository = transcript_segment_repository
        self.extraction_run_repository = extraction_run_repository
        self.action_repository = action_repository
        self.decision_repository = decision_repository
        self.risk_repository = risk_repository
        self.question_repository = question_repository
        self.topic_repository = topic_repository
        self.evidence_repository = evidence_repository
        self.summary_repository = summary_repository
        self.artifact_repository = artifact_repository
        self.export_run_repository = export_run_repository

    def create_export(
        self,
        *,
        meeting_id: int,
        export_profile: ExportProfile,
        format: ExportFormat,
        options: ExportOptions,
        output_directory: str | None = None,
    ) -> dict[str, Any]:
        source = self._build_source(meeting_id=meeting_id, options=options)
        descriptor = self._build_descriptor(
            meeting_id=meeting_id,
            meeting_title=source["meeting"]["title"],
            meeting_project=source["meeting"].get("project"),
            meeting_date=source["meeting"].get("meeting_date"),
            export_profile=export_profile,
            format=format,
            output_directory=output_directory,
        )
        export_run_id = self.export_run_repository.create(
            meeting_id=meeting_id,
            export_profile=export_profile,
            format=format,
            options_json=options.to_json(),
            file_path=str(descriptor.file_path),
        )
        try:
            self._generate_export(descriptor=descriptor, source=source, options=options)
            self.export_run_repository.finalize_success(export_run_id, file_path=str(descriptor.file_path))
        except Exception as exc:  # noqa: BLE001
            self.export_run_repository.finalize_failure(export_run_id, error_message=str(exc))
            raise
        return self.export_run_repository.get(export_run_id) or {}

    def list_exports(self, meeting_id: int) -> list[dict[str, Any]]:
        return self.export_run_repository.list_for_meeting(meeting_id)

    def get_export(self, export_run_id: int) -> dict[str, Any] | None:
        return self.export_run_repository.get(export_run_id)

    def open_export_folder(self, export_run_id: int) -> dict[str, Any]:
        export_run = self.export_run_repository.get(export_run_id)
        if not export_run:
            raise ValueError("Export run not found")
        export_path = Path(export_run["file_path"])
        folder = export_path.parent
        ensure_directory(folder)
        if hasattr(os, "startfile"):
            os.startfile(str(folder))  # type: ignore[attr-defined]
        return {"status": "opened", "folder_path": str(folder)}

    def _build_source(self, *, meeting_id: int, options: ExportOptions) -> dict[str, Any]:
        meeting = self.meeting_repository.get(meeting_id)
        if not meeting:
            raise ValueError("Meeting not found")
        source_file = self.source_file_repository.get_for_meeting(meeting_id)
        preprocessing_run = self.preprocessing_run_repository.get_latest_for_meeting(meeting_id)
        transcription_run = self.transcription_run_repository.get_latest_for_meeting(meeting_id)
        extraction_run = self.extraction_run_repository.get_latest_for_meeting(meeting_id)
        if not source_file:
            raise ValueError("Meeting source file not found")

        merged_segments = []
        if transcription_run and transcription_run["status"] in {"completed", "completed_with_failures", "recovered"}:
            merged_segments = self.transcript_segment_repository.list_for_run(
                int(transcription_run["id"]),
                "merged",
                include_excluded=False,
            )

        extraction_payload = {
            "run": extraction_run,
            "summary": None,
            "actions": [],
            "decisions": [],
            "risks": [],
            "questions": [],
            "topics": [],
            "evidence_links": [],
        }
        if extraction_run and extraction_run["status"] == "completed":
            extraction_payload["summary"] = self.summary_repository.get_for_run(int(extraction_run["id"]))
            raw_actions = self.action_repository.list_for_run(int(extraction_run["id"]))
            raw_decisions = self.decision_repository.list_for_run(int(extraction_run["id"]))
            raw_risks = self.risk_repository.list_for_run(int(extraction_run["id"]))
            raw_questions = self.question_repository.list_for_run(int(extraction_run["id"]))
            raw_topics = self.topic_repository.list_for_run(int(extraction_run["id"]))
            evidence_links = self.evidence_repository.list_for_run(int(extraction_run["id"]))
            grouped = self._group_evidence(evidence_links)
            extraction_payload["actions"] = self._attach_evidence(raw_actions, grouped, "action", reviewed_only=options.reviewed_only)
            extraction_payload["decisions"] = self._attach_evidence(raw_decisions, grouped, "decision", reviewed_only=options.reviewed_only)
            extraction_payload["risks"] = self._attach_evidence(raw_risks, grouped, "risk", reviewed_only=options.reviewed_only)
            extraction_payload["questions"] = self._attach_evidence(raw_questions, grouped, "question", reviewed_only=options.reviewed_only)
            extraction_payload["topics"] = self._attach_evidence(raw_topics, grouped, "topic", reviewed_only=options.reviewed_only)
            extraction_payload["evidence_links"] = evidence_links

        if not merged_segments and options.include_transcript_appendix:
            raise ValueError("Merged transcript is not available for transcript appendix export")
        return {
            "meeting": meeting,
            "source_file": source_file,
            "preprocessing_run": preprocessing_run,
            "transcription_run": transcription_run,
            "merged_segments": merged_segments,
            "extraction": extraction_payload,
            "artifacts": self.artifact_repository.list_for_meeting(meeting_id),
            "generated_at": utc_now_iso(),
        }

    def _generate_export(self, *, descriptor: ExportDescriptor, source: dict[str, Any], options: ExportOptions) -> None:
        if descriptor.export_profile == "formal_minutes_pack":
            summary = source["extraction"]["summary"]
            if not summary:
                raise ValueError("Formal minutes export requires a completed extraction summary")
            if descriptor.format == "docx":
                self._generate_docx(descriptor.file_path, source, options)
                return
            if descriptor.format == "pdf":
                self._generate_pdf(descriptor.file_path, source, options)
                return
        if descriptor.export_profile == "action_register":
            if not source["extraction"]["actions"]:
                raise ValueError("Action register export requires persisted extracted actions")
            if descriptor.format == "csv":
                self._generate_csv(descriptor.file_path, source, options)
                return
            if descriptor.format == "xlsx":
                self._generate_xlsx(descriptor.file_path, source, options)
                return
        if descriptor.export_profile == "full_archive" and descriptor.format == "json":
            self._generate_json(descriptor.file_path, source, options)
            return
        if descriptor.export_profile == "transcript_export" and descriptor.format == "txt":
            if not source["merged_segments"]:
                raise ValueError("Transcript export requires a merged transcript")
            self._generate_txt(descriptor.file_path, source)
            return
        raise ValueError(f"Unsupported export combination: {descriptor.export_profile} / {descriptor.format}")

    def _build_descriptor(
        self,
        *,
        meeting_id: int,
        meeting_title: str,
        meeting_project: str | None,
        meeting_date: str | None,
        export_profile: ExportProfile,
        format: ExportFormat,
        output_directory: str | None,
    ) -> ExportDescriptor:
        if output_directory:
            export_root = ensure_directory(Path(output_directory))
        else:
            export_root = ensure_directory(config.exports_root / f"meeting_{meeting_id}")

        base_label = _export_filename_label(
            project=meeting_project,
            meeting_date=meeting_date,
            fallback_title=meeting_title,
        )
        existing_exports = self.export_run_repository.list_for_meeting(meeting_id)
        sequence = f"{len(existing_exports) + 1:02d}"
        file_name = f"{base_label} - {sequence}.{format}"
        return ExportDescriptor(
            export_profile=export_profile,
            format=format,
            file_path=export_root / file_name,
            display_name=file_name,
        )

    def _generate_docx(self, path: Path, source: dict[str, Any], options: ExportOptions) -> None:
        document = Document()
        self._configure_docx_styles(document)
        section = document.sections[0]
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

        meeting = source["meeting"]
        summary = source["extraction"]["summary"]
        self._add_docx_title_block(document, source)

        self._add_docx_section_band(document, "Executive Summary")
        document.add_paragraph(summary["summary_text"] or "No executive summary available.", style="AE2 Body")

        self._add_docx_section_band(document, "Formal Minutes")
        self._add_docx_minutes(document, summary["minutes_text"] or "No formal minutes available.")

        self._add_docx_items(document, "Decisions", source["extraction"]["decisions"], options)
        self._add_docx_items(document, "Risks And Issues", source["extraction"]["risks"], options)
        self._add_docx_items(document, "Open Questions", source["extraction"]["questions"], options)

        self._add_docx_section_band(document, "Action Items")
        actions_table = document.add_table(rows=1, cols=4)
        actions_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        actions_table.autofit = False
        self._apply_docx_table_widths(actions_table, _docx_action_column_widths())
        headers = ["Action", "Owner", "Due Date", "Priority"]
        header_row = actions_table.rows[0]
        self._shade_docx_row(header_row, "0F172A")
        for index, header in enumerate(headers):
            cell = header_row.cells[index]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.style = document.styles["AE2 Table Header"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.add_run(header)
        for action in source["extraction"]["actions"]:
            row = actions_table.add_row().cells
            row[0].text = action["text"]
            row[1].text = action.get("owner") or "-"
            row[2].text = _format_display_date(action.get("due_date"))
            row[3].text = _humanize_value(action.get("priority"))
            self._apply_docx_table_widths(actions_table, _docx_action_column_widths())
        self._style_docx_actions_table(document, actions_table)

        if options.include_evidence_appendix:
            document.add_section(WD_SECTION.NEW_PAGE)
            document.add_paragraph("Evidence Appendix", style="AE2 Heading")
            for section_name, items in (
                ("Actions", source["extraction"]["actions"]),
                ("Decisions", source["extraction"]["decisions"]),
                ("Risks", source["extraction"]["risks"]),
                ("Questions", source["extraction"]["questions"]),
            ):
                document.add_paragraph(section_name, style="AE2 Subheading")
                if not items:
                    document.add_paragraph("No items.", style="AE2 Body")
                for item in items:
                    document.add_paragraph(item["text"], style="AE2 Bullet")
                    for evidence in item.get("evidence") or []:
                        document.add_paragraph(_evidence_line(evidence, include_confidence=options.include_confidence_flags), style="AE2 Evidence")

        if options.include_transcript_appendix and source["merged_segments"]:
            document.add_section(WD_SECTION.NEW_PAGE)
            document.add_paragraph("Transcript Appendix", style="AE2 Heading")
            for segment in source["merged_segments"]:
                document.add_paragraph(_transcript_line(segment), style="AE2 Transcript")

        ensure_directory(path.parent)
        document.save(path)

    def _generate_pdf(self, path: Path, source: dict[str, Any], options: ExportOptions) -> None:
        ensure_directory(path.parent)
        styles = self._pdf_styles()
        doc = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=16 * mm, leftMargin=16 * mm, topMargin=14 * mm, bottomMargin=14 * mm)
        story: list[Any] = []
        meeting = source["meeting"]
        summary = source["extraction"]["summary"]

        story.append(Paragraph(meeting["title"], styles["title"]))
        story.append(Spacer(1, 6))
        metadata = [
            ["Meeting date", _format_display_date(meeting.get("meeting_date"))],
            ["Project", meeting.get("project") or "-"],
        ]
        meta_table = Table(metadata, colWidths=[35 * mm, 130 * mm])
        meta_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#111827")),
            ("TEXTCOLOR", (0, 0), (-1, -1), colors.whitesmoke),
            ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#334155")),
            ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
            ("FONTSIZE", (0, 0), (-1, -1), 8.5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
        ]))
        story.append(meta_table)
        story.append(Spacer(1, 10))

        story.extend(self._pdf_section("Executive Summary", summary["summary_text"], styles))
        story.extend(self._pdf_section("Formal Minutes", summary["minutes_text"], styles))
        story.extend(self._pdf_insight_section("Decisions", source["extraction"]["decisions"], styles, options))
        story.extend(self._pdf_insight_section("Risks And Issues", source["extraction"]["risks"], styles, options))
        story.extend(self._pdf_insight_section("Open Questions", source["extraction"]["questions"], styles, options))

        story.append(Paragraph("Action Items", styles["heading"]))
        action_rows = [["Action", "Owner", "Due Date", "Priority"]]
        for action in source["extraction"]["actions"]:
            action_rows.append([
                action["text"],
                action.get("owner") or "-",
                _format_display_date(action.get("due_date")),
                _humanize_value(action.get("priority")),
            ])
        action_table = Table(action_rows, colWidths=[95 * mm, 32 * mm, 28 * mm, 24 * mm], repeatRows=1)
        action_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0f172a")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
            ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#334155")),
            ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
            ("FONTSIZE", (0, 0), (-1, -1), 7.8),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(action_table)

        if options.include_evidence_appendix:
            story.append(Spacer(1, 10))
            story.append(Paragraph("Evidence Appendix", styles["heading"]))
            for label, items in (
                ("Actions", source["extraction"]["actions"]),
                ("Decisions", source["extraction"]["decisions"]),
                ("Risks", source["extraction"]["risks"]),
                ("Questions", source["extraction"]["questions"]),
            ):
                story.append(Paragraph(label, styles["subheading"]))
                for item in items:
                    story.append(Paragraph(item["text"], styles["body"]))
                    for evidence in item.get("evidence") or []:
                        story.append(Paragraph(_evidence_line(evidence, include_confidence=options.include_confidence_flags), styles["evidence"]))
                    story.append(Spacer(1, 3))

        if options.include_transcript_appendix and source["merged_segments"]:
            story.append(Spacer(1, 10))
            story.append(Paragraph("Transcript Appendix", styles["heading"]))
            for segment in source["merged_segments"]:
                story.append(Paragraph(_transcript_line(segment), styles["transcript"]))

        doc.build(story)

    def _generate_csv(self, path: Path, source: dict[str, Any], options: ExportOptions) -> None:
        ensure_directory(path.parent)
        with path.open("w", newline="", encoding="utf-8") as handle:
            writer = csv.writer(handle)
            writer.writerow(["action_text", "owner", "due_date", "review_status", "explicit_or_inferred", "evidence_timestamps", "confidence"])
            for action in source["extraction"]["actions"]:
                writer.writerow([
                    action["text"],
                    action.get("owner") or "",
                    action.get("due_date") or "",
                    action.get("review_status") or "",
                    action.get("explicit_or_inferred") or "",
                    "; ".join(_evidence_timestamps(action.get("evidence") or [])),
                    _confidence_text(action.get("confidence"), options),
                ])

    def _generate_xlsx(self, path: Path, source: dict[str, Any], options: ExportOptions) -> None:
        ensure_directory(path.parent)
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Actions"
        headers = ["Action Text", "Owner", "Due Date", "Review Status", "Explicit Or Inferred", "Evidence Timestamps", "Confidence"]
        sheet.append(headers)
        for cell in sheet[1]:
            font = copy(cell.font)
            font.bold = True
            cell.font = font
        for action in source["extraction"]["actions"]:
            sheet.append([
                action["text"],
                action.get("owner") or "",
                action.get("due_date") or "",
                action.get("review_status") or "",
                action.get("explicit_or_inferred") or "",
                "; ".join(_evidence_timestamps(action.get("evidence") or [])),
                action.get("confidence") if options.include_confidence_flags else "",
            ])
        widths = [44, 20, 16, 16, 18, 28, 12]
        for index, width in enumerate(widths, start=1):
            sheet.column_dimensions[chr(64 + index)].width = width
        workbook.save(path)

    def _generate_json(self, path: Path, source: dict[str, Any], options: ExportOptions) -> None:
        ensure_directory(path.parent)
        payload = {
            "meeting": source["meeting"],
            "source_file": source["source_file"],
            "preprocessing_run": source["preprocessing_run"],
            "transcription_run": source["transcription_run"],
            "merged_transcript": source["merged_segments"],
            "extraction": source["extraction"],
            "artifacts": source["artifacts"],
            "export_generated_at": source["generated_at"],
            "export_options": options.to_json(),
        }
        path.write_text(json.dumps(payload, indent=2), encoding="utf-8")

    def _generate_txt(self, path: Path, source: dict[str, Any]) -> None:
        ensure_directory(path.parent)
        lines = [_transcript_line(segment) for segment in source["merged_segments"]]
        path.write_text("\n".join(lines), encoding="utf-8")

    @staticmethod
    def _group_evidence(evidence_links: list[dict[str, Any]]) -> dict[tuple[str, int], list[dict[str, Any]]]:
        grouped: dict[tuple[str, int], list[dict[str, Any]]] = {}
        for entry in evidence_links:
            grouped.setdefault((entry["entity_type"], entry["entity_id"]), []).append(entry)
        return grouped

    @staticmethod
    def _attach_evidence(items: list[dict[str, Any]], grouped: dict[tuple[str, int], list[dict[str, Any]]], entity_type: str, *, reviewed_only: bool) -> list[dict[str, Any]]:
        enriched = []
        for item in items:
            if reviewed_only and item.get("review_status") != "accepted":
                continue
            evidence = sorted(grouped.get((entity_type, item["id"]), []), key=lambda row: (row["start_ms"], row["end_ms"], row["id"]))
            enriched.append({**item, "evidence": evidence})
        return enriched

    @staticmethod
    def _configure_docx_styles(document: Document) -> None:
        styles = document.styles
        if "AE2 Title" not in styles:
            title_style = styles.add_style("AE2 Title", WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = "Aptos Display"
            title_style.font.size = Pt(22)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(15, 23, 42)
            title_style.paragraph_format.space_after = Pt(3)
        if "AE2 Eyebrow" not in styles:
            eyebrow_style = styles.add_style("AE2 Eyebrow", WD_STYLE_TYPE.PARAGRAPH)
            eyebrow_style.font.name = "Aptos"
            eyebrow_style.font.size = Pt(8.5)
            eyebrow_style.font.bold = True
            eyebrow_style.font.color.rgb = RGBColor(8, 145, 178)
            eyebrow_style.paragraph_format.space_after = Pt(2)
        if "AE2 Heading" not in styles:
            heading_style = styles.add_style("AE2 Heading", WD_STYLE_TYPE.PARAGRAPH)
            heading_style.font.name = "Aptos Display"
            heading_style.font.size = Pt(13.5)
            heading_style.font.bold = True
            heading_style.font.color.rgb = RGBColor(15, 23, 42)
            heading_style.paragraph_format.space_before = Pt(10)
            heading_style.paragraph_format.space_after = Pt(4)
        if "AE2 Section Band" not in styles:
            section_band_style = styles.add_style("AE2 Section Band", WD_STYLE_TYPE.PARAGRAPH)
            section_band_style.font.name = "Aptos Display"
            section_band_style.font.size = Pt(13)
            section_band_style.font.bold = True
            section_band_style.font.color.rgb = RGBColor(255, 255, 255)
            section_band_style.paragraph_format.space_before = Pt(8)
            section_band_style.paragraph_format.space_after = Pt(0)
        if "AE2 Subheading" not in styles:
            subheading_style = styles.add_style("AE2 Subheading", WD_STYLE_TYPE.PARAGRAPH)
            subheading_style.font.name = "Aptos"
            subheading_style.font.size = Pt(10.5)
            subheading_style.font.bold = True
            subheading_style.font.color.rgb = RGBColor(51, 65, 85)
            subheading_style.paragraph_format.space_before = Pt(6)
            subheading_style.paragraph_format.space_after = Pt(2)
        if "AE2 Body" not in styles:
            body_style = styles.add_style("AE2 Body", WD_STYLE_TYPE.PARAGRAPH)
            body_style.font.name = "Aptos"
            body_style.font.size = Pt(10.25)
            body_style.font.color.rgb = RGBColor(30, 41, 59)
            body_style.paragraph_format.line_spacing = 1.15
            body_style.paragraph_format.space_after = Pt(5)
        if "AE2 Minutes Heading" not in styles:
            minutes_heading_style = styles.add_style("AE2 Minutes Heading", WD_STYLE_TYPE.PARAGRAPH)
            minutes_heading_style.font.name = "Aptos"
            minutes_heading_style.font.size = Pt(9.5)
            minutes_heading_style.font.bold = True
            minutes_heading_style.font.color.rgb = RGBColor(15, 23, 42)
            minutes_heading_style.paragraph_format.space_before = Pt(3)
            minutes_heading_style.paragraph_format.space_after = Pt(1)
        if "AE2 Bullet" not in styles:
            bullet_style = styles.add_style("AE2 Bullet", WD_STYLE_TYPE.PARAGRAPH)
            bullet_style.base_style = styles["List Bullet"]
            bullet_style.font.name = "Aptos"
            bullet_style.font.size = Pt(10)
            bullet_style.font.color.rgb = RGBColor(30, 41, 59)
            bullet_style.paragraph_format.space_before = Pt(0)
            bullet_style.paragraph_format.space_after = Pt(2)
        if "AE2 Evidence" not in styles:
            evidence_style = styles.add_style("AE2 Evidence", WD_STYLE_TYPE.PARAGRAPH)
            evidence_style.font.name = "Aptos"
            evidence_style.font.size = Pt(8.75)
            evidence_style.font.color.rgb = RGBColor(71, 85, 105)
            evidence_style.paragraph_format.left_indent = Inches(0.2)
            evidence_style.paragraph_format.space_after = Pt(2)
        if "AE2 Transcript" not in styles:
            transcript_style = styles.add_style("AE2 Transcript", WD_STYLE_TYPE.PARAGRAPH)
            transcript_style.font.name = "Consolas"
            transcript_style.font.size = Pt(9)
            transcript_style.paragraph_format.space_after = Pt(2)
        if "AE2 Meta Label" not in styles:
            meta_label_style = styles.add_style("AE2 Meta Label", WD_STYLE_TYPE.PARAGRAPH)
            meta_label_style.font.name = "Aptos"
            meta_label_style.font.size = Pt(8)
            meta_label_style.font.bold = True
            meta_label_style.font.color.rgb = RGBColor(71, 85, 105)
        if "AE2 Meta Value" not in styles:
            meta_value_style = styles.add_style("AE2 Meta Value", WD_STYLE_TYPE.PARAGRAPH)
            meta_value_style.font.name = "Aptos"
            meta_value_style.font.size = Pt(10)
            meta_value_style.font.color.rgb = RGBColor(15, 23, 42)
        if "AE2 Table Header" not in styles:
            table_header_style = styles.add_style("AE2 Table Header", WD_STYLE_TYPE.PARAGRAPH)
            table_header_style.font.name = "Aptos"
            table_header_style.font.size = Pt(8.5)
            table_header_style.font.bold = True
            table_header_style.font.color.rgb = RGBColor(255, 255, 255)
        if "AE2 Table Body" not in styles:
            table_body_style = styles.add_style("AE2 Table Body", WD_STYLE_TYPE.PARAGRAPH)
            table_body_style.font.name = "Aptos"
            table_body_style.font.size = Pt(9)
            table_body_style.font.color.rgb = RGBColor(30, 41, 59)

    def _add_docx_title_block(self, document: Document, source: dict[str, Any]) -> None:
        meeting = source["meeting"]
        metadata_rows = [
            ("Meeting date", _format_display_date(meeting.get("meeting_date"))),
            ("Project", meeting.get("project") or "-"),
        ]
        document.add_paragraph(meeting["title"], style="AE2 Title")

        metadata_table = document.add_table(rows=1, cols=2)
        metadata_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        metadata_table.autofit = False
        metadata_table.rows[0].cells[0].width = Inches(3.2)
        metadata_table.rows[0].cells[1].width = Inches(3.2)
        for index, (label, value) in enumerate(metadata_rows):
            cell = metadata_table.cell(index // 2, index % 2)
            self._shade_docx_cell(cell, "F8FAFC")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            label_paragraph = cell.paragraphs[0]
            label_paragraph.style = document.styles["AE2 Meta Label"]
            label_paragraph.add_run(label.upper())
            value_paragraph = cell.add_paragraph(str(value), style="AE2 Meta Value")
            value_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        spacer = document.add_paragraph("")
        spacer.paragraph_format.space_after = Pt(1)

        self._add_docx_people_section(document, "Attendees", meeting.get("attendees") or [])
        self._add_docx_people_section(document, "Circulation", meeting.get("circulation") or [])

    def _style_docx_actions_table(self, document: Document, table) -> None:
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                for paragraph in cell.paragraphs:
                    paragraph.style = document.styles["AE2 Table Header" if row_index == 0 else "AE2 Table Body"]
            if row_index == 0:
                continue
            if row_index % 2 == 1:
                self._shade_docx_row(row, "F8FAFC")

    @staticmethod
    def _apply_docx_table_widths(table, widths: tuple[float, ...]) -> None:
        for column, width in zip(table.columns, widths, strict=True):
            column.width = Inches(width)
        for row in table.rows:
            for cell, width in zip(row.cells, widths, strict=True):
                cell.width = Inches(width)

    @staticmethod
    def _shade_docx_row(row, fill: str) -> None:
        for cell in row.cells:
            ExportService._shade_docx_cell(cell, fill)

    @staticmethod
    def _shade_docx_cell(cell, fill: str) -> None:
        cell._tc.get_or_add_tcPr().append(parse_xml(rf'<w:shd {nsdecls("w")} w:fill="{fill}"/>'))

    @staticmethod
    def _add_docx_items(document: Document, heading: str, items: list[dict[str, Any]], options: ExportOptions) -> None:
        ExportService._add_docx_section_band(document, heading)
        if not items:
            document.add_paragraph("No items available.", style="AE2 Body")
            return
        for item in items:
            paragraph = document.add_paragraph(style="AE2 Bullet")
            lead = paragraph.add_run(item["text"])
            lead.bold = True
            if options.include_confidence_flags:
                paragraph.add_run(f"  [{_confidence_text(item.get('confidence'), options)} confidence]")

    @staticmethod
    def _add_docx_minutes(document: Document, minutes_text: str) -> None:
        blocks = _parse_minutes_blocks(minutes_text)
        if not blocks:
            document.add_paragraph("No formal minutes available.", style="AE2 Body")
            return
        for block in blocks:
            if block["kind"] == "heading":
                ExportService._add_docx_minutes_heading_block(document, str(block["text"]).upper())
            elif block["kind"] == "bullet":
                document.add_paragraph(str(block["text"]), style="AE2 Bullet")
            else:
                document.add_paragraph(str(block["text"]), style="AE2 Body")

    @staticmethod
    def _add_docx_minutes_heading_block(document: Document, text: str) -> None:
        paragraph = document.add_paragraph(style="AE2 Minutes Heading")
        run = paragraph.add_run(text)
        run.bold = True

    @staticmethod
    def _add_docx_section_band(document: Document, text: str) -> None:
        table = document.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        cell = table.rows[0].cells[0]
        cell.width = Inches(6.5)
        ExportService._shade_docx_cell(cell, "0F3A52")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.style = document.styles["AE2 Section Band"]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.add_run(text)

    @staticmethod
    def _add_docx_people_section(document: Document, heading: str, people: list[str]) -> None:
        if not people:
            return
        document.add_paragraph(heading, style="AE2 Subheading")
        for person in people:
            document.add_paragraph(person, style="AE2 Bullet")

    @staticmethod
    def _pdf_styles() -> dict[str, ParagraphStyle]:
        base = getSampleStyleSheet()
        return {
            "eyebrow": ParagraphStyle("AE2Eyebrow", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=8, textColor=colors.HexColor("#0891b2"), spaceAfter=4),
            "title": ParagraphStyle("AE2Title", parent=base["Heading1"], fontName="Helvetica-Bold", fontSize=18, leading=22, textColor=colors.HexColor("#0f172a"), spaceAfter=10),
            "heading": ParagraphStyle("AE2Heading", parent=base["Heading2"], fontName="Helvetica-Bold", fontSize=12, leading=15, textColor=colors.HexColor("#0f172a"), spaceBefore=6, spaceAfter=4),
            "subheading": ParagraphStyle("AE2Subheading", parent=base["Heading3"], fontName="Helvetica-Bold", fontSize=10, leading=12, textColor=colors.HexColor("#0f172a"), spaceBefore=4, spaceAfter=3),
            "body": ParagraphStyle("AE2Body", parent=base["BodyText"], fontName="Helvetica", fontSize=9.5, leading=13, spaceAfter=5),
            "evidence": ParagraphStyle("AE2Evidence", parent=base["BodyText"], fontName="Helvetica", fontSize=8.2, leading=11, leftIndent=10),
            "transcript": ParagraphStyle("AE2Transcript", parent=base["Code"], fontName="Courier", fontSize=7.8, leading=10.5, spaceAfter=2),
        }

    @staticmethod
    def _pdf_section(title: str, text: str, styles: dict[str, ParagraphStyle]) -> list[Any]:
        return [Paragraph(title, styles["heading"]), Paragraph((text or "No content available.").replace("\n", "<br/>"), styles["body"]), Spacer(1, 5)]

    def _pdf_insight_section(self, title: str, items: list[dict[str, Any]], styles: dict[str, ParagraphStyle], options: ExportOptions) -> list[Any]:
        flow: list[Any] = [Paragraph(title, styles["heading"])]
        if not items:
            flow.append(Paragraph("No items available.", styles["body"]))
            flow.append(Spacer(1, 4))
            return flow
        for item in items:
            text = item["text"]
            if options.include_confidence_flags:
                text = f"{text} ({_confidence_text(item.get('confidence'), options)})"
            flow.append(Paragraph(f"• {text}", styles["body"]))
        flow.append(Spacer(1, 4))
        return flow


def _slugify(value: str) -> str:
    cleaned = "".join(character.lower() if character.isalnum() else "_" for character in value.strip())
    collapsed = "_".join(segment for segment in cleaned.split("_") if segment)
    return collapsed[:48] or "meeting"


def _sanitize_filename_label(value: str) -> str:
    sanitized = re.sub(r'[<>:"/\\|?*]+', " ", str(value or "").strip())
    sanitized = re.sub(r"\s+", " ", sanitized).strip(" .-")
    return sanitized


def _export_filename_label(*, project: str | None, meeting_date: str | None, fallback_title: str) -> str:
    label = _sanitize_filename_label(project or fallback_title or "Meeting")
    date_label = _sanitize_filename_label(meeting_date or utc_now_iso()[:10])
    return f"{label} - {date_label}"


def _humanize_value(value: Any) -> str:
    if value is None:
        return "-"
    return str(value).replace("_", " ").strip().title()


def _docx_action_column_widths() -> tuple[float, float, float, float]:
    # Keep the action register within the printable width of the default DOCX layout.
    return (3.55, 1.45, 0.95, 0.55)


def _format_display_date(value: Any) -> str:
    if not value:
        return "-"
    text = str(value).strip()
    for pattern in ("%Y-%m-%d", "%Y-%m-%dT%H:%M:%S", "%Y-%m-%dT%H:%M:%S.%f", "%Y-%m-%dT%H:%M:%S%z", "%Y-%m-%dT%H:%M:%S.%f%z"):
        try:
            from datetime import datetime

            return datetime.strptime(text, pattern).strftime("%d-%B-%y")
        except ValueError:
            continue
    if re.fullmatch(r"\d{4}-\d{2}-\d{2}", text):
        from datetime import datetime

        return datetime.strptime(text, "%Y-%m-%d").strftime("%d-%B-%y")
    return text


def _parse_minutes_blocks(minutes_text: str) -> list[dict[str, str]]:
    blocks: list[dict[str, str]] = []
    normalized_text = _split_inline_minutes_headings(str(minutes_text or ""))
    for raw_line in normalized_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if _looks_like_bullet(line):
            blocks.append({"kind": "bullet", "text": _strip_bullet_prefix(line)})
            continue
        if _looks_like_minutes_heading(line):
            blocks.append({"kind": "heading", "text": line.rstrip(":")})
            continue
        blocks.append({"kind": "paragraph", "text": line})
    return blocks


def _split_inline_minutes_headings(minutes_text: str) -> str:
    text = str(minutes_text or "").strip()
    if not text:
        return ""

    known_headings = [
        "KEY DISCUSSION TOPICS",
        "DISCUSSION TOPICS",
        "DECISIONS MADE",
        "DECISIONS",
        "RISKS & ISSUES",
        "RISKS AND ISSUES",
        "RISKS",
        "ISSUES",
        "ACTION ITEMS",
        "ACTIONS",
        "OPEN QUESTIONS",
        "QUESTIONS",
        "NEXT STEPS",
        "FOLLOW UPS",
    ]

    normalized = text.replace("\r\n", "\n")
    heading_pattern = re.compile(
        r"(?<!\n)\s*(?P<heading>" + "|".join(re.escape(heading) for heading in sorted(known_headings, key=len, reverse=True)) + r"):\s*",
        re.IGNORECASE,
    )

    def replace_heading(match: re.Match[str]) -> str:
        heading = str(match.group("heading")).upper()
        return f"\n{heading}:\n" if match.start() > 0 else f"{heading}:\n"

    normalized = heading_pattern.sub(replace_heading, normalized)
    normalized = re.sub(r"\s+-\s+", "\n- ", normalized)
    return normalized.strip()


def _looks_like_minutes_heading(line: str) -> bool:
    normalized = line.strip().rstrip(":")
    if not normalized:
        return False
    if len(normalized) > 60:
        return False
    words = normalized.split()
    if not words:
        return False
    if line.endswith(":"):
        return True
    alpha_words = [word for word in words if any(char.isalpha() for char in word)]
    if not alpha_words:
        return False
    uppercase_words = [word for word in alpha_words if word == word.upper()]
    return len(alpha_words) <= 6 and len(uppercase_words) == len(alpha_words)


def _looks_like_bullet(line: str) -> bool:
    stripped = line.lstrip()
    if stripped.startswith(("-", "*", "•")):
        return True
    if len(stripped) >= 3 and stripped[0].isdigit() and stripped[1] in {".", ")"} and stripped[2] == " ":
        return True
    return False


def _strip_bullet_prefix(line: str) -> str:
    stripped = line.lstrip()
    if stripped.startswith(("-", "*", "•")):
        return stripped[1:].strip()
    if len(stripped) >= 3 and stripped[0].isdigit() and stripped[1] in {".", ")"} and stripped[2] == " ":
        return stripped[3:].strip()
    return stripped


def _format_timestamp(ms: int) -> str:
    total_seconds = max(0, ms // 1000)
    hours, remainder = divmod(total_seconds, 3600)
    minutes, seconds = divmod(remainder, 60)
    return f"{hours:02d}:{minutes:02d}:{seconds:02d}"


def _evidence_timestamps(evidence: list[dict[str, Any]]) -> list[str]:
    return [f"{_format_timestamp(item['start_ms'])}-{_format_timestamp(item['end_ms'])}" for item in evidence]


def _confidence_text(confidence: Any, options: ExportOptions) -> str:
    if not options.include_confidence_flags:
        return "-"
    if confidence is None:
        return "n/a"
    return f"{round(float(confidence) * 100)}%"


def _evidence_line(evidence: dict[str, Any], *, include_confidence: bool) -> str:
    confidence = ""
    if include_confidence and evidence.get("confidence") is not None:
        confidence = f" | confidence {round(float(evidence['confidence']) * 100)}%"
    speaker = f" | {evidence['speaker_label']}" if evidence.get("speaker_label") else ""
    quote = f" | {evidence.get('quote_snippet')}" if evidence.get("quote_snippet") else ""
    return f"{_format_timestamp(int(evidence['start_ms']))}-{_format_timestamp(int(evidence['end_ms']))}{speaker}{confidence}{quote}"


def _transcript_line(segment: dict[str, Any]) -> str:
    speaker = segment.get("speaker_name") or segment.get("speaker_label") or "Speaker"
    return f"[{_format_timestamp(int(segment['start_ms_in_meeting']))} - {_format_timestamp(int(segment['end_ms_in_meeting']))}] {speaker}: {segment['text']}"
